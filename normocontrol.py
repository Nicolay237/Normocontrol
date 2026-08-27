#!/usr/bin/env python3
# -*- coding: utf-8 -*-


import sys
import os
import re
import math
import argparse
import statistics
import html as html_lib
from collections import Counter
from dataclasses import dataclass, field
from datetime import datetime

NORMS = {
    "margins_cm": {"left": 3.0, "right": 1.0, "top": 2.0, "bottom": 2.0},
    "margins_tolerance_cm": 0.1,
    "margins_tolerance_cm_pdf": 0.3,
    "margins_tolerance_cm_pdf_bottom": 0.8,
    "font_name": "Times New Roman",
    "font_size_pt": 14,
    "font_size_tolerance_pt": 0.5,
    "line_spacing": 1.5,
    "line_spacing_tolerance": 0.05,
    "line_spacing_tolerance_pdf": 0.2,
    "line_height_factor": 1.15,
    "first_line_indent_cm": 1.25,
    "first_line_indent_tolerance_cm": 0.1,
    "first_line_indent_tolerance_cm_pdf": 0.3,
    "justify_min_ratio": 0.6,
    "heading_font_size_pt": 14,
    "heading_bold": True,
    # По ГОСТ 7.32-2017 в оглавлении жирным выделяются только заголовки
    # верхнего уровня (ВВЕДЕНИЕ, РАЗДЕЛ N, ЗАКЛЮЧЕНИЕ и т.п.),
    # пункты/подпункты (1.1, 1.1.1) — обычным начертанием.
    "toc_top_level_bold": True,
    "toc_sub_level_bold": False,
}

HEADING_STYLE_PREFIXES = ("Heading", "Заголовок", "Title", "Название")
PT_TO_CM = 2.54 / 72.0


@dataclass
class Issue:
    location: str
    category: str
    message: str


@dataclass
class Report:
    issues: list = field(default_factory=list)
    notes: list = field(default_factory=list)

    def add(self, location, category, message):
        self.issues.append(Issue(location, category, message))

    def note(self, text):
        self.notes.append(text)

    def is_clean(self):
        return len(self.issues) == 0


def iter_all_paragraphs(document):
    """
    Обходит ВСЕ абзацы верхнего уровня в теле документа, включая те,
    что спрятаны внутри content control (<w:sdt>).

    Word оборачивает автособираемое оглавление именно в такой sdt-блок,
    поэтому document.paragraphs (стандартный способ python-docx) его
    просто не видит — из-за этого чекер был "слеп" ко всему оглавлению
    целиком (не только к жирному шрифту в нём).

    В таблицы не спускается — там абзацы обходятся отдельно, как и раньше
    (см. doc.tables в check_docx).
    """
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph

    body = document.element.body

    def walk(parent_element):
        for child in parent_element:
            tag = child.tag.split('}')[-1]
            if tag == 'p':
                yield Paragraph(child, document)
            elif tag == 'sdt':
                sdt_content = child.find(qn('w:sdtContent'))
                if sdt_content is not None:
                    yield from walk(sdt_content)

    yield from walk(body)


def get_all_runs(paragraph):
    """
    Возвращает ВСЕ runs абзаца, включая те, что лежат внутри <w:hyperlink>.

    Строки оглавления — это внутренние гиперссылки на разделы, а
    paragraph.runs (стандартный способ python-docx) runs внутри
    <w:hyperlink> не отдаёт, поэтому проверка жирности по paragraph.runs
    для записей оглавления всегда возвращала пустой список.
    """
    from docx.oxml.ns import qn
    from docx.text.run import Run

    return [Run(r, paragraph) for r in paragraph._p.findall('.//' + qn('w:r'))]


def toc_style_level(p):
    """
    Определяет уровень записи оглавления (1, 2, 3...) по имени стиля
    абзаца. Word называет такие стили 'TOC 1'/'TOC 2'/... или
    'Оглавление 1'/'Оглавление 2'/... в зависимости от локали шаблона.
    Возвращает None, если абзац не является записью оглавления по стилю.
    """
    name = (p.style.name if p.style else "").strip().lower()
    m = re.match(r"(?:toc|оглавлени[ея]|содержани[ея])\s*(\d+)", name)
    return int(m.group(1)) if m else None


def estimate_pagination(doc, paragraphs):
    #Грубая оценка того, на какой странице окажется каждый абзац
    section = doc.sections[0]
    content_w = section.page_width.pt - section.left_margin.pt - section.right_margin.pt
    content_h = section.page_height.pt - section.top_margin.pt - section.bottom_margin.pt

    mapping = {}
    page, used_lines = 1, 0.0
    page_counts = Counter()

    for idx, p in enumerate(paragraphs, start=1):
        text = p.text.strip()
        size = NORMS["font_size_pt"]
        for run in get_all_runs(p):
            if run.font.size:
                size = run.font.size.pt
                break
        spacing = p.paragraph_format.line_spacing
        try:
            spacing = float(spacing) if spacing else NORMS["line_spacing"]
        except (TypeError, ValueError):
            spacing = NORMS["line_spacing"]

        line_height = size * NORMS["line_height_factor"] * spacing
        chars_per_line = max(1, content_w / (size * 0.5))
        lines_per_page = max(1, content_h / line_height)
        n_lines = max(1, math.ceil(len(text) / chars_per_line)) if text else 1

        if used_lines > 0 and used_lines + n_lines > lines_per_page:
            page += 1
            used_lines = 0
        used_lines += n_lines
        page_counts[page] += 1
        mapping[idx] = (page, page_counts[page])

    return mapping


def check_docx(path, report):
    from docx import Document
    from docx.shared import Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    MATH_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"
    FORMULA_FONTS = {"cambria math", "symbol", "mt extra", "euclid math one"}
    LIST_MARKERS = re.compile(r"^\s*([-–—•*]|\d+[.)])\s+")

    def emu_to_cm(v):
        return round(Emu(v).cm, 3) if v is not None else None

    def pt(v):
        return round(v.pt, 2) if v is not None else None

    def preview(p, n=60):
        t = p.text.strip()
        return t[:n] + "..." if len(t) > n else (t or "(пустой абзац)")

    def is_heading(p):
        name = p.style.name if p.style else ""
        return any(name.startswith(x) for x in HEADING_STYLE_PREFIXES)

    def is_formula(p):
        return p._element.find(f".//{MATH_NS}oMath") is not None

    def is_toc_entry(p):
        name = (p.style.name if p.style else "").lower()
        if "toc" in name or "оглавлен" in name or "содержан" in name:
            return True
        text = p.text
        # запись оглавления обычно заканчивается табуляцией/точками-лидерами и номером страницы
        return bool(re.search(r"[\t.]{2,}\s*\d+\s*$", text)) or bool(re.search(r"\t\s*\d+\s*$", text))

    def is_list_item(p):
        name = (p.style.name if p.style else "").lower()
        return "list" in name or bool(LIST_MARKERS.match(p.text))

    def run_font(run, default):
        if run.font.name:
            return run.font.name
        rpr = run._element.find(
            './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts'
        )
        if rpr is not None:
            f = rpr.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii')
            if f:
                return f
        return default

    def default_font(doc):
        try:
            return doc.styles['Normal'].font.name or "Times New Roman"
        except Exception:
            return "Times New Roman"

    def effective(p, attr):
        pf = p.paragraph_format
        if getattr(pf, attr) is not None:
            return getattr(pf, attr)
        style = p.style
        while style is not None:
            spf = style.paragraph_format
            if spf and getattr(spf, attr) is not None:
                return getattr(spf, attr)
            style = style.base_style
        return None

    doc = Document(path)
    dfont = default_font(doc)
    all_paragraphs = list(iter_all_paragraphs(doc))
    page_map = estimate_pagination(doc, all_paragraphs)
    report.note(
        "Номера страниц — приблизительная оценка по объёму текста и параметрам "
        "страницы, реальная разбивка Word может отличаться (таблицы, изображения, разрывы)."
    )

    def locate(idx, p):
        page, n = page_map.get(idx, (None, None))
        prefix = f"Страница {page}, абзац {n}" if page else f"Абзац №{idx}"
        return f"{prefix} («{preview(p)}»)"

    def locate_heading(idx, p):
        page, n = page_map.get(idx, (None, None))
        prefix = f"Страница {page}, заголовок в абзаце {n}" if page else f"Заголовок в абзаце №{idx}"
        return f"{prefix} («{preview(p)}»)"

    def locate_toc(idx, p):
        page, n = page_map.get(idx, (None, None))
        prefix = f"Страница {page}, оглавление, абзац {n}" if page else f"Оглавление, абзац №{idx}"
        return f"{prefix} («{preview(p)}»)"

    for i, s in enumerate(doc.sections, start=1):
        vals = {
            "left": emu_to_cm(s.left_margin), "right": emu_to_cm(s.right_margin),
            "top": emu_to_cm(s.top_margin), "bottom": emu_to_cm(s.bottom_margin),
        }
        tol = NORMS["margins_tolerance_cm"]
        for side, expected in NORMS["margins_cm"].items():
            actual = vals[side]
            if actual is not None and abs(actual - expected) > tol:
                report.add(f"Раздел документа №{i}", "Поля страницы",
                           f"поле '{side}' = {actual} см, ожидается {expected} см (допуск ±{tol} см)")

    for idx, p in enumerate(all_paragraphs, start=1):
        if not p.text.strip() or is_heading(p) or is_toc_entry(p):
            continue
        loc = locate(idx, p)
        formula = is_formula(p)
        list_item = is_list_item(p)

        for run in get_all_runs(p):
            if not run.text.strip():
                continue
            fname = run_font(run, dfont)
            if fname and fname != NORMS["font_name"] and fname.lower() not in FORMULA_FONTS and not formula:
                report.add(loc, "Шрифт", f"использован шрифт '{fname}', ожидается '{NORMS['font_name']}'")
            size = pt(run.font.size)
            if size is not None and abs(size - NORMS["font_size_pt"]) > NORMS["font_size_tolerance_pt"] and not formula:
                report.add(loc, "Кегль шрифта", f"размер шрифта {size} pt, ожидается {NORMS['font_size_pt']} pt")

        spacing = effective(p, "line_spacing")
        if spacing is not None:
            try:
                v = float(spacing)
                if abs(v - NORMS["line_spacing"]) > NORMS["line_spacing_tolerance"]:
                    report.add(loc, "Межстрочный интервал", f"интервал {v}, ожидается {NORMS['line_spacing']}")
            except (TypeError, ValueError):
                pass

        if not list_item and not formula:
            alignment = effective(p, "alignment")
            if alignment is not None and alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
                report.add(loc, "Выравнивание", f"выравнивание '{alignment}', ожидается 'JUSTIFY' (по ширине)")

            indent = effective(p, "first_line_indent")
            indent_cm = emu_to_cm(indent) if indent is not None else None
            expected = NORMS["first_line_indent_cm"]
            tol = NORMS["first_line_indent_tolerance_cm"]
            if indent_cm is None or abs(indent_cm - expected) > tol:
                report.add(loc, "Красная строка", f"отступ первой строки {indent_cm or 0} см, ожидается {expected} см")

    for idx, p in enumerate(all_paragraphs, start=1):
        if not is_heading(p) or not p.text.strip() or is_toc_entry(p):
            continue
        loc = locate_heading(idx, p)
        runs = [r for r in get_all_runs(p) if r.text.strip()]
        if NORMS["heading_bold"] and runs and not all(r.bold for r in runs):
            report.add(loc, "Оформление заголовка", "заголовок должен быть выделен жирным (bold)")
        if p.text.strip().endswith("."):
            report.add(loc, "Оформление заголовка", "в конце заголовка не должно быть точки")
        for run in runs:
            size = pt(run.font.size)
            if size is not None and size != NORMS["heading_font_size_pt"]:
                report.add(loc, "Оформление заголовка",
                           f"размер шрифта заголовка {size} pt, ожидается {NORMS['heading_font_size_pt']} pt")

    # Оформление оглавления: заголовки верхнего уровня должны быть жирными,
    # пункты/подпункты — нет. Раньше это было невозможно проверить в принципе:
    # doc.paragraphs не видел записи оглавления (см. iter_all_paragraphs),
    # а paragraph.runs не видел runs внутри гиперссылок (см. get_all_runs).
    for idx, p in enumerate(all_paragraphs, start=1):
        level = toc_style_level(p)
        if level is None or not p.text.strip():
            continue
        # Номер страницы — это отдельный run (результат поля PAGEREF) и
        # обычно НЕ наследует жирность заголовка, даже когда сам заголовок
        # жирный. Это не ошибка оформления, поэтому цифры-номера страниц
        # исключаются из проверки, чтобы не было ложных срабатываний.
        runs = [
            r for r in get_all_runs(p)
            if r.text.strip() and not re.fullmatch(r"\d+", r.text.strip())
        ]
        if not runs:
            continue
        loc = locate_toc(idx, p)
        is_bold = all(r.bold for r in runs)
        if level == 1 and NORMS["toc_top_level_bold"] and not is_bold:
            report.add(loc, "Оформление оглавления",
                       "заголовок раздела в оглавлении должен быть выделен жирным (bold)")
        elif level > 1 and not NORMS["toc_sub_level_bold"] and is_bold:
            report.add(loc, "Оформление оглавления",
                       "пункт оглавления не должен быть жирным (жирным выделяются только "
                       "заголовки разделов верхнего уровня)")

    for ti, table in enumerate(doc.tables, start=1):
        for ri, row in enumerate(table.rows, start=1):
            for ci, cell in enumerate(row.cells, start=1):
                for p in cell.paragraphs:
                    if not p.text.strip():
                        continue
                    for run in p.runs:
                        if not run.text.strip():
                            continue
                        fname = run_font(run, dfont)
                        if fname and fname != NORMS["font_name"] and fname.lower() not in FORMULA_FONTS:
                            report.add(f"Таблица №{ti}, строка {ri}, столбец {ci}", "Шрифт (таблица)",
                                       f"использован шрифт '{fname}', ожидается '{NORMS['font_name']}'")

    check_typography(
        [(locate(i, p), p.text) for i, p in enumerate(all_paragraphs, start=1)
         if not is_toc_entry(p)],
        report,
    )



def _cluster_lines(words, y_tol=2.5):
    lines, current, top = [], [], None
    for w in sorted(words, key=lambda w: (round(w["top"], 1), w["x0"])):
        if top is None or abs(w["top"] - top) <= y_tol:
            current.append(w)
            top = w["top"] if top is None else top
        else:
            lines.append(current)
            current, top = [w], w["top"]
    if current:
        lines.append(current)
    return lines


def check_pdf(path, report):
    import pdfplumber

    report.note(
        "PDF не хранит структуру абзацев и стилей — показатели ниже вычислены "
        "по координатам текста и являются оценочными."
    )

    all_x0, all_x1, font_counter, page_texts = [], [], Counter(), []

    with pdfplumber.open(path) as pdf:
        if not pdf.pages:
            report.add("Документ", "Общее", "в PDF не найдено ни одной страницы")
            return

        for p_idx, page in enumerate(pdf.pages, start=1):
            words = page.extract_words(extra_attrs=["fontname", "size"])
            if not words:
                report.note(f"Страница {p_idx}: текстовый слой не найден (возможно, скан) — пропущена.")
                continue

            chars = page.chars
            page_w, page_h = page.width, page.height
            x0s = [c["x0"] for c in chars]
            x1s = [c["x1"] for c in chars]
            tops = [c["top"] for c in chars]
            bottoms = [c["bottom"] for c in chars]

            measured = {
                "left": round(min(x0s) * PT_TO_CM, 2),
                "right": round((page_w - max(x1s)) * PT_TO_CM, 2),
                "top": round(min(tops) * PT_TO_CM, 2),
                "bottom": round((page_h - max(bottoms)) * PT_TO_CM, 2),
            }
            skip_bottom = measured["bottom"] > page_h * PT_TO_CM * 0.35
            if skip_bottom:
                report.note(f"Страница {p_idx}: текст не доходит до низа листа — нижнее поле не проверялось.")

            tol = NORMS["margins_tolerance_cm_pdf"]
            for side, expected in NORMS["margins_cm"].items():
                if side == "bottom" and skip_bottom:
                    continue
                side_tol = NORMS["margins_tolerance_cm_pdf_bottom"] if side == "bottom" else tol
                if abs(measured[side] - expected) > side_tol:
                    report.add(f"Страница {p_idx}", "Поля страницы",
                               f"оценочное поле '{side}' ≈ {measured[side]} см, ожидается {expected} см (допуск ±{side_tol} см)")

            for w in words:
                font_counter[(w.get("fontname"), round(w.get("size", 0), 1))] += 1

            page_lines = []
            for line in _cluster_lines(words):
                lw = sorted(line, key=lambda w: w["x0"])
                x0, x1, top = min(w["x0"] for w in lw), max(w["x1"] for w in lw), min(w["top"] for w in lw)
                text = " ".join(w["text"] for w in lw)
                sizes = [w.get("size") for w in lw if w.get("size")]
                size = round(statistics.median(sizes), 1) if sizes else None
                page_lines.append({"x0": x0, "x1": x1, "top": top, "text": text, "size": size})
                all_x0.append(round(x0, 1))
                all_x1.append(round(x1, 1))
                page_texts.append((f"Страница {p_idx}, строка «{text[:50]}»", text))

            if page_lines:
                body_x0 = Counter(round(l["x0"], 0) for l in page_lines).most_common(1)[0][0]
                prev_top = None
                for line in page_lines:
                    continuation = abs(round(line["x0"], 0) - body_x0) <= 3
                    if prev_top is not None and line["size"] and continuation:
                        gap = line["top"] - prev_top
                        expected_gap = NORMS["line_spacing"] * line["size"] * NORMS["line_height_factor"]
                        if expected_gap > 0 and 0 < gap < expected_gap * 2.5:
                            if abs(gap - expected_gap) / expected_gap > NORMS["line_spacing_tolerance_pdf"]:
                                t = line["text"]
                                report.add(f"Страница {p_idx}, строка «{t[:50]}»", "Межстрочный интервал (оценка)",
                                           f"промежуток ≈ {round(gap,1)} pt, ожидается ≈ {round(expected_gap,1)} pt "
                                           f"(интервал {NORMS['line_spacing']})")
                    prev_top = line["top"]

        if font_counter:
            total = sum(font_counter.values())
            for (fontname, size), count in font_counter.items():
                if count / total < 0.01:
                    continue
                name_ok = NORMS["font_name"].split()[0].lower() in (fontname or "").lower()
                size_ok = abs(size - NORMS["font_size_pt"]) <= NORMS["font_size_tolerance_pt"]
                if not name_ok or not size_ok:
                    reasons = []
                    if not name_ok:
                        reasons.append(f"шрифт '{fontname}' (ожидается '{NORMS['font_name']}')")
                    if not size_ok:
                        reasons.append(f"кегль {size} pt (ожидается {NORMS['font_size_pt']} pt)")
                    report.add(f"Документ (встречается {count} раз, слов)", "Шрифт / кегль",
                               "обнаружено использование: " + ", ".join(reasons))

        if all_x0:
            counts = Counter(all_x0)
            body_x0 = counts.most_common(1)[0][0]
            expected_pt = NORMS["first_line_indent_cm"] / PT_TO_CM
            tol_pt = NORMS["first_line_indent_tolerance_cm_pdf"] / PT_TO_CM
            candidates = [x for x in counts if x > body_x0 + 3]
            if candidates:
                typical = statistics.mode([round(x - body_x0, 1) for x in candidates])
                if abs(typical - expected_pt) > tol_pt:
                    report.add("Документ (оценка по всем страницам)", "Красная строка (оценка)",
                               f"типичный отступ первой строки ≈ {round(typical * PT_TO_CM, 2)} см, "
                               f"ожидается {NORMS['first_line_indent_cm']} см")
            else:
                report.note("Не обнаружен характерный отступ красной строки — похоже, абзацы не имеют отступа.")

        if all_x1:
            body_right = Counter(all_x1).most_common(1)[0][0]
            ratio = sum(1 for x in all_x1 if abs(x - body_right) <= 5) / len(all_x1)
            if ratio < NORMS["justify_min_ratio"]:
                report.add("Документ (оценка по всем страницам)", "Выравнивание (оценка)",
                           f"только {round(ratio*100)}% строк доходят до правого края — похоже на выравнивание "
                           f"по левому краю (ожидается по ширине)")

    check_typography(page_texts, report)


def check_typography(labeled_texts, report):
    prev_empty = False
    for loc, text in labeled_texts:
        stripped = (text or "").strip()
        if not stripped:
            if prev_empty:
                report.add(loc, "Пустые строки", "два и более пустых абзаца подряд")
            prev_empty = True
            continue
        prev_empty = False
        if "  " in text:
            report.add(loc, "Пробелы", "обнаружен двойной (или более) пробел")
        for punct in [" ,", " .", " ;", " :", " !", " ?"]:
            if punct in text:
                report.add(loc, "Пробелы", f"лишний пробел перед знаком препинания '{punct.strip()}'")
        if text.startswith(" ") or text.startswith("\t"):
            report.add(loc, "Пробелы", "текст начинается с пробела/табуляции вместо красной строки")


def run_normocontrol(path):
    report = Report()
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        check_docx(path, report)
    elif ext == ".pdf":
        check_pdf(path, report)
    else:
        raise ValueError(f"Неподдерживаемый формат файла: '{ext}'. Ожидается .docx или .pdf")
    return report


REPORT_CSS = """
.report-card {
  max-width: 860px; margin: 0 auto;
  background: #fffdf8;
  border: 1px solid var(--line);
  box-shadow: 0 2px 24px rgba(0,0,0,0.08);
  padding: 48px 56px;
  position: relative;
  overflow: hidden;
}
.report-card::before {
  content: ""; position: absolute; left: 0; top: 0; bottom: 0; width: 6px;
  background: linear-gradient(var(--gold), var(--red));
}
.report-card header { display: flex; justify-content: space-between; align-items: flex-start; gap: 24px; }
.report-card h1 { font-size: 26px; margin: 0 0 6px; letter-spacing: 0.3px; }
.report-card .meta { color: var(--ink-soft); font-size: 14px; font-family: ui-monospace, SFMono-Regular, Menlo, monospace; }
.report-card .stamp {
  flex-shrink: 0; border: 3px solid; border-radius: 50%;
  width: 132px; height: 132px; display: flex; align-items: center; justify-content: center;
  text-align: center; font-weight: bold; font-size: 13px; letter-spacing: 0.5px;
  transform: rotate(-9deg); font-family: "Courier New", ui-monospace, monospace;
  padding: 10px;
}
.report-card .stamp-bad { color: var(--red); border-color: var(--red); }
.report-card .stamp-ok { color: var(--green); border-color: var(--green); }
.report-card hr { border: none; border-top: 1px solid var(--line); margin: 28px 0; }
.report-card .notes { color: var(--ink-soft); font-size: 14px; margin: 0 0 24px; padding-left: 18px; }
.report-card .clean-msg { font-size: 17px; color: var(--green); }
.report-card details.category { border: 1px solid var(--line); border-radius: 4px; margin-bottom: 14px; background: #fffefb; }
.report-card summary {
  cursor: pointer; padding: 12px 18px; font-size: 15px; font-weight: bold;
  display: flex; justify-content: space-between; align-items: center;
  list-style: none;
}
.report-card summary::-webkit-details-marker { display: none; }
.report-card .cat-count {
  background: var(--red); color: #fff; border-radius: 12px;
  font-size: 12px; padding: 2px 9px; font-family: ui-monospace, monospace;
}
.report-card table { width: 100%; border-collapse: collapse; font-size: 14px; }
.report-card td { padding: 8px 18px; border-top: 1px solid var(--line); vertical-align: top; }
.report-card td.loc {
  width: 38%; color: var(--ink-soft); font-family: ui-monospace, SFMono-Regular, Menlo, monospace;
  font-size: 12.5px;
}
.report-card footer { margin-top: 32px; color: var(--ink-soft); font-size: 12px; }
@media print { .report-card { box-shadow: none; border: none; } }
"""

REPORT_CSS_VARS = """
:root {
  --paper: #f6f2e8; --ink: #23241f; --ink-soft: #5b5a52;
  --red: #a3323d; --green: #3f6b46; --gold: #b9964f; --line: #ddd5c0;
}
"""


def build_report_fragment(report, source_path):
    #Возвращает HTML-фрагмент отчёта (без <html>/<head>)
    e = html_lib.escape
    by_category = {}
    for issue in report.issues:
        by_category.setdefault(issue.category, []).append(issue)

    clean = report.is_clean()
    total = len(report.issues)
    stamp_text = "СООТВЕТСТВУЕТ&nbsp;НОРМАМ" if clean else f"{total}&nbsp;ЗАМЕЧАН{'ИЕ' if total == 1 else 'ИЙ'}"
    stamp_class = "stamp-ok" if clean else "stamp-bad"

    notes_html = "".join(f'<li>{e(n)}</li>' for n in report.notes)
    notes_block = f'<ul class="notes">{notes_html}</ul>' if notes_html else ""

    sections_html = ""
    for category, issues in by_category.items():
        rows = "".join(
            f'<tr><td class="loc">{e(i.location)}</td><td>{e(i.message)}</td></tr>'
            for i in issues
        )
        sections_html += f"""
        <details class="category" open>
          <summary><span class="cat-name">{e(category)}</span><span class="cat-count">{len(issues)}</span></summary>
          <table>{rows}</table>
        </details>"""

    body_html = notes_block + (
        '<p class="clean-msg">Замечаний не обнаружено. Документ соответствует заданным нормам.</p>'
        if clean else sections_html
    )

    return f"""<div class="report-card">
    <header>
      <div>
        <h1>Отчёт нормоконтроля</h1>
        <div class="meta">{e(os.path.basename(source_path))} · {datetime.now().strftime('%d.%m.%Y %H:%M')}</div>
      </div>
      <div class="stamp {stamp_class}">{stamp_text}</div>
    </header>
    <hr>
    {body_html}
    <footer>Нормы: ГОСТ 7.32-2017 (типовые требования для ВКР) · сгенерировано normocontrol.py</footer>
  </div>"""


def build_html_report(report, source_path):
    e = html_lib.escape
    fragment = build_report_fragment(report, source_path)
    return f"""<!DOCTYPE html>
<html lang="ru">
<head>
<meta charset="utf-8">
<title>Нормоконтроль — {e(os.path.basename(source_path))}</title>
<style>
{REPORT_CSS_VARS}
body {{ margin: 0; padding: 40px 20px; background: var(--paper); color: var(--ink);
  font-family: Georgia, "Times New Roman", serif; line-height: 1.5; }}
{REPORT_CSS}
</style>
</head>
<body>
{fragment}
</body>
</html>"""


def format_text_report(report, source_path):
    lines = ["=" * 70, "ОТЧЁТ НОРМОКОНТРОЛЯ", f"Файл: {source_path}", "=" * 70]
    if report.notes:
        lines.append("")
        lines += [f"ⓘ {n}" for n in report.notes]
    if report.is_clean():
        lines.append("\nЗамечаний не обнаружено. Документ соответствует заданным нормам.")
        return "\n".join(lines)
    lines.append(f"\nВсего замечаний: {len(report.issues)}\n")
    by_category = {}
    for issue in report.issues:
        by_category.setdefault(issue.category, []).append(issue)
    for category, issues in by_category.items():
        lines.append(f"--- {category} ({len(issues)}) ---")
        lines += [f"  [{i.location}] {i.message}" for i in issues]
        lines.append("")
    return "\n".join(lines)


def main():
    parser = argparse.ArgumentParser(description="Проверка .docx/.pdf на нормоконтроль")
    parser.add_argument("doc_path", help="путь к .docx или .pdf файлу")
    parser.add_argument("--html", help="путь для HTML-отчёта (по умолчанию: <файл>_report.html)")
    parser.add_argument("--report", help="путь для текстового отчёта (опционально)")
    parser.add_argument("--no-html", action="store_true", help="не создавать HTML-отчёт")
    args = parser.parse_args()

    try:
        report = run_normocontrol(args.doc_path)
    except Exception as e:
        print(f"Ошибка при обработке файла: {e}", file=sys.stderr)
        sys.exit(1)

    print(format_text_report(report, args.doc_path))

    if args.report:
        with open(args.report, "w", encoding="utf-8") as f:
            f.write(format_text_report(report, args.doc_path))
        print(f"\nТекстовый отчёт сохранён в: {args.report}")

    if not args.no_html:
        html_path = args.html or os.path.splitext(args.doc_path)[0] + "_report.html"
        with open(html_path, "w", encoding="utf-8") as f:
            f.write(build_html_report(report, args.doc_path))
        print(f"HTML-отчёт сохранён в: {html_path}")


if __name__ == "__main__":
    main()